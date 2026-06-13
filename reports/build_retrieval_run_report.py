from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("reports/bao_cao_retrieval_can_lam_de_chay_duoc.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if name == "Heading 2" else 16)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Báo cáo trạng thái Retrieval: cần làm gì để chạy được")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Branch: feature/retrievel | Phạm vi: Phase 0 và Phase 1 | Ngày: 13/06/2026")
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. Kết luận nhanh", level=1)
    doc.add_paragraph(
        "Role Retrieval đã có module code thật cho baseline visual search: nhận query text, "
        "encode bằng OpenCLIP, tìm trong FAISS, map kết quả qua frame_map và trả về result chuẩn. "
        "Module đã pass unit test và compile check."
    )
    doc.add_paragraph(
        "Tuy nhiên hệ thống chưa chạy live end-to-end trên repo hiện tại vì thiếu artifact dữ liệu "
        "do pipeline Indexing/Metadata tạo ra: FAISS index và frame_map."
    )

    doc.add_heading("2. Trạng thái hiện tại", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Hạng mục", "Trạng thái", "Ghi chú"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows = [
        (
            "Code Retrieval",
            "Đã có",
            "VisualSearchEngine, RetrievalManager, API wrapper, response model.",
        ),
        (
            "Unit test",
            "Đã pass",
            "Test fake encoder/searcher kiểm tra mapping FAISS id sang metadata result.",
        ),
        (
            "OpenCLIP",
            "Đã tích hợp lazy-load",
            "Sẽ load thư viện/model khi chạy search thật.",
        ),
        (
            "FAISS index",
            "Thiếu",
            "Chưa có data/indexes/openclip_vit_b16_flat_ip.faiss trong repo.",
        ),
        (
            "frame_map",
            "Thiếu",
            "Chưa có data/metadata/openclip_vit_b16_frame_map.json trong repo.",
        ),
        (
            "Benchmark live",
            "Chưa chạy được",
            "Bị chặn bởi thiếu artifact runtime.",
        ),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
    set_table_width(table, [1.55, 1.45, 3.5])

    doc.add_heading("3. Vì sao OpenCLIP vẫn chưa đủ để chạy?", level=1)
    doc.add_paragraph(
        "OpenCLIP chỉ tạo embedding cho query text. Retrieval vẫn cần một kho vector ảnh đã được "
        "encode trước từ keyframes để so khớp. Kho đó nằm trong FAISS index. Sau khi FAISS trả về "
        "các chỉ số như 0, 1, 2, hệ thống cần frame_map để biết chỉ số đó thuộc video nào, frame nào, "
        "timestamp nào và ảnh keyframe nằm ở đâu."
    )
    add_bullet(doc, "OpenCLIP = biến query text thành vector.")
    add_bullet(doc, "FAISS index = kho vector của toàn bộ keyframes đã encode.")
    add_bullet(doc, "frame_map.json = bản đồ từ faiss_index về video_id, frame_id, timestamp, keyframe_path.")

    doc.add_heading("4. Cần làm gì để chạy được live?", level=1)
    add_number(doc, "P2 chuẩn bị folder keyframes thật trong data/keyframes/<video_id>.")
    add_number(doc, "P2 chạy pipeline OpenCLIP để tạo embedding .npy và metadata embedding.")
    add_number(doc, "P2 build FAISS index từ embedding đã tạo.")
    add_number(doc, "P3/P2 bảo đảm frame_map.json khớp đúng số vector trong FAISS.")
    add_number(doc, "P4 chạy search_visual với query mẫu và ghi latency/top-k result.")
    add_number(doc, "P1/P4 bổ sung query mẫu và ground truth để tính Recall@K, MRR.")

    doc.add_heading("5. Lệnh chạy đề xuất", level=1)
    doc.add_paragraph("Tạo embedding cho một folder keyframes:")
    add_code_block(
        doc,
        [
            "python backend/app/services/indexing/run_keyframe_openclip_pipeline.py \\",
            "  --keyframe-dir data/keyframes/<video_id> \\",
            "  --output-root data",
        ],
    )
    doc.add_paragraph("Build FAISS index:")
    add_code_block(
        doc,
        [
            "python backend/app/services/indexing/build_faiss_index.py \\",
            "  --embeddings-glob \"data/embeddings/openclip_vit_b16_*.npy\" \\",
            "  --index-path data/indexes/openclip_vit_b16_flat_ip.faiss \\",
            "  --frame-map-path data/metadata/openclip_vit_b16_frame_map.json",
        ],
    )
    doc.add_paragraph("Test Retrieval thật sau khi có artifact:")
    add_code_block(
        doc,
        [
            "python -c \"from backend.app.services.retrieval.retrieval_manager import search_visual; "
            "print(search_visual('a man cooking in a kitchen', 20).to_dict())\"",
        ],
    )

    doc.add_heading("6. Điều kiện hoàn thành Phase 1 cho Retrieval", level=1)
    add_bullet(doc, "Có FAISS index và frame_map thật trong data/.")
    add_bullet(doc, "Query text trả về top-k keyframes có score, video_id, frame_id, timestamp, keyframe_path.")
    add_bullet(doc, "Chạy được ít nhất 20 query mẫu.")
    add_bullet(doc, "Có báo cáo latency trung bình và lỗi thực tế.")
    add_bullet(doc, "Nếu có ground truth, báo cáo Recall@10, Recall@50 và MRR.")

    doc.add_heading("7. Kết luận", level=1)
    doc.add_paragraph(
        "Retrieval hiện đã sẵn sàng để nối với artifact của Indexing/Metadata. Việc tiếp theo không phải "
        "viết thêm nhiều code Retrieval, mà là tạo đúng FAISS index và frame_map từ dữ liệu thật. Sau khi "
        "có hai file này, P4 có thể chạy live benchmark và chuyển sang hybrid/rerank."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
